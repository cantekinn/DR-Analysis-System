# -*- coding: utf-8 -*-
"""
EEM0458 Görüntü İşleme Projesi - Otomatik Word Rapor Üreticisi
Diyabetik Retinopati Analizi için Derin Öğrenme Tabanlı Çoklu Görev Sistemi
Yazar: Can Tekin (cantekinn)
Bahar 2025-2026

Kullanım:
    pip install python-docx
    python generate_report.py

Çıktı: Rapor.docx (aynı klasörde)

Figürlerin C:\\Users\\Can Tekin\\Desktop\\lna\\DR_Project\\figures\\ altında olması gerekir.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
FIG_DIR = BASE_DIR / 'figures'
OUTPUT = BASE_DIR / 'Rapor.docx'

doc = Document()


def set_default_font(doc, name='Times New Roman', size=12):
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        style.element.append(rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


set_default_font(doc, 'Times New Roman', 12)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def add_para(text='', bold=False, size=12, align='left', space_after=6,
             line_spacing=1.5, indent_first=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    alignments = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                  'center': WD_ALIGN_PARAGRAPH.CENTER,
                  'right': WD_ALIGN_PARAGRAPH.RIGHT,
                  'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_heading_main(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'


def add_heading_sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def add_figure(fname, caption, width_cm=14):
    fpath = FIG_DIR / fname
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    if fpath.exists():
        run.add_picture(str(fpath), width=Cm(width_cm))
    else:
        run.add_text(f"[GÖRSEL EKSİK: {fname}]")
        run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = 1.15
    cr = cap.add_run(caption)
    cr.bold = True
    cr.font.size = Pt(11)
    cr.font.name = 'Times New Roman'


def add_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    if col_widths_cm:
        for row in table.rows:
            for i, c in enumerate(row.cells):
                if i < len(col_widths_cm):
                    c.width = Cm(col_widths_cm[i])
    return table


def page_break():
    doc.add_page_break()


# ====================== KAPAK SAYFASI ======================
add_para('BURSA TEKNİK ÜNİVERSİTESİ', bold=True, size=14, align='center',
         space_after=6, line_spacing=1.15)
add_para('MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ', bold=True, size=14,
         align='center', space_after=24, line_spacing=1.15)

add_para('[BTÜ LOGOSU BURAYA]', size=11, align='center', space_after=12)
add_para('', space_after=24)

add_para('DİYABETİK RETİNOPATİ ANALİZİ İÇİN DERİN ÖĞRENME TABANLI '
         'ÇOKLU GÖREV SİSTEMİ', bold=True, size=14, align='center',
         space_after=48, line_spacing=1.5)

add_para('', space_after=48)
add_para('EEM0458 Görüntü İşleme Projesi', bold=True, size=12, align='center',
         space_after=12)
add_para('Can Tekin', bold=True, size=12, align='center', space_after=24)
add_para('', space_after=48)
add_para('Elektrik Elektronik Mühendisliği Bölümü', bold=True, size=12,
         align='center', space_after=200)

add_para('2025/2026 Bahar Dönemi', bold=True, size=12, align='center')

page_break()


# ====================== İÇİNDEKİLER ======================
add_para('İÇİNDEKİLER', bold=True, size=14, align='left', space_after=24)
add_para('                                                                              Sayfa',
         bold=True, size=11, align='left', space_after=12)

icindekiler = [
    ('İÇİNDEKİLER', 'ii'),
    ('KISALTMALAR', 'iii'),
    ('SEMBOLLER', 'iv'),
    ('ŞEKİL LİSTESİ', 'v'),
    ('ÖZET', 'vi'),
    ('1. GİRİŞ', '7'),
    ('2. LİTERATÜR TARAMASI', '8'),
    ('3. MATERYAL VE YÖNTEM', '10'),
    ('   3.1 Veri Setleri', '10'),
    ('   3.2 Görüntü Ön İşleme', '11'),
    ('   3.3 Veri Bölümleme', '12'),
    ('   3.4 Sınıflandırma Modeli', '12'),
    ('   3.5 Segmentasyon Modeli', '13'),
    ('   3.6 Açıklanabilirlik (Grad-CAM)', '14'),
    ('   3.7 Eğitim Detayları', '14'),
    ('4. BULGULAR', '15'),
    ('   4.1 Sınıflandırma Sonuçları', '15'),
    ('   4.2 Domain Adaptation Sonuçları', '17'),
    ('   4.3 Segmentasyon Sonuçları', '18'),
    ('   4.4 Grad-CAM Analizi', '20'),
    ('5. TARTIŞMA VE SINIRLILIKLAR', '21'),
    ('6. SONUÇ', '23'),
    ('KAYNAKLAR', '24'),
]
for item, sayfa in icindekiler:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), alignment=2, leader=1)
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if not item.startswith(' '):
        run.bold = True
    p.add_run('\t').font.size = Pt(11)
    r2 = p.add_run(sayfa)
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

page_break()


# ====================== KISALTMALAR ======================
add_para('KISALTMALAR', bold=True, size=14, align='left', space_after=24)

kisaltmalar = [
    ('AI', 'Artificial Intelligence (Yapay Zeka)'),
    ('AMD', 'Age-related Macular Degeneration'),
    ('APTOS', 'Asia Pacific Tele-Ophthalmology Society'),
    ('BCE', 'Binary Cross-Entropy'),
    ('CLAHE', 'Contrast Limited Adaptive Histogram Equalization'),
    ('CNN', 'Convolutional Neural Network (Evrişimli Sinir Ağı)'),
    ('DR', 'Diabetic Retinopathy (Diyabetik Retinopati)'),
    ('EX', 'Hard Exudate (Sert Eksuda)'),
    ('Grad-CAM', 'Gradient-weighted Class Activation Mapping'),
    ('HE', 'Haemorrhage (Kanama)'),
    ('IDRiD', 'Indian Diabetic Retinopathy Image Dataset'),
    ('IoU', 'Intersection over Union'),
    ('MA', 'Microaneurysm (Mikroanevrizma)'),
    ('PDR', 'Proliferative Diabetic Retinopathy'),
    ('RGB', 'Red, Green, Blue'),
    ('SE', 'Soft Exudate (Yumuşak Eksuda)'),
    ('U-Net', 'U-shaped Convolutional Network'),
    ('VRAM', 'Video Random Access Memory'),
]
for kis, acl in kisaltmalar:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f'{kis:12s}')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(f': {acl}')
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

page_break()


# ====================== SEMBOLLER ======================
add_para('SEMBOLLER', bold=True, size=14, align='left', space_after=24)
semboller = [
    ('m', 'Görüntü Genişliği'),
    ('n', 'Görüntü Yüksekliği'),
    ('h(x,y)', 'Piksel Yoğunluk Değeri'),
    ('p', 'Olasılık Vektörü'),
    ('y', 'Gerçek Etiket'),
    ('L', 'Kayıp Fonksiyonu'),
    ('η', 'Öğrenme Oranı (Learning Rate)'),
    ('κ', 'Cohen Quadratic Kappa Katsayısı'),
    ('Dice', '2|A ∩ B| / (|A| + |B|)'),
    ('IoU', '|A ∩ B| / |A ∪ B|'),
]
for sem, acl in semboller:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f'{sem:12s}')
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(11)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(f': {acl}')
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

page_break()


# ====================== ŞEKİL LİSTESİ ======================
add_para('ŞEKİL LİSTESİ', bold=True, size=14, align='left', space_after=24)

sekiller = [
    ('Şekil 1.', 'Veri setlerinin sınıf bazlı dağılımı', '10'),
    ('Şekil 2.', 'Fundus görüntü ön-işleme pipeline (Orijinal → Crop → CLAHE → Final)', '11'),
    ('Şekil 3.', 'Veri artırma (augmentation) örnekleri', '12'),
    ('Şekil 4.', 'Sınıflandırma modeli eğitim eğrileri (EfficientNet-B3, 20 epoch)', '15'),
    ('Şekil 5.', 'Karışıklık matrisleri: APTOS validasyon ve IDRiD external test', '16'),
    ('Şekil 6.', 'IDRiD-B üzerinde fine-tune sonrası karışıklık matrisleri', '17'),
    ('Şekil 7.', 'Fine-tune etkisi: APTOS Val ve IDRiD Test Kappa skoru', '17'),
    ('Şekil 8.', 'Segmentasyon modeli eğitim eğrileri (U-Net, 60 epoch)', '18'),
    ('Şekil 9.', 'Lezyon bazında final Dice ve IoU skorları', '19'),
    ('Şekil 10.', 'Test seti örnekleri üzerinde tahmin maskeleri (GT vs Pred)', '19'),
    ('Şekil 11.', '5 sınıf için Grad-CAM++ açıklanabilirlik ısı haritaları', '20'),
]
for kod, baslik, sayfa in sekiller:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), alignment=2, leader=1)
    r1 = p.add_run(f'{kod:10s}')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(baslik)
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'
    p.add_run('\t').font.size = Pt(11)
    r3 = p.add_run(sayfa)
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

page_break()


# ====================== ÖZET ======================
add_para('DİYABETİK RETİNOPATİ ANALİZİ İÇİN DERİN ÖĞRENME TABANLI '
         'ÇOKLU GÖREV SİSTEMİ', bold=True, size=12, align='center',
         space_after=18, line_spacing=1.5)
add_para('ÖZET', bold=True, size=12, align='center', space_after=18)

ozet = (
    "Diyabetik retinopati (DR), diyabet hastalığının neden olduğu ve "
    "çalışma yaşındaki yetişkinler arasında görme kaybının başlıca "
    "sebeplerinden biri olan mikrovasküler bir komplikasyondur. Hastalığın "
    "erken evrede tespiti görme kaybını önlemede kritik bir role sahiptir. "
    "Bu projede, fundus (göz dibi) görüntülerinden DR şiddet sınıflandırması "
    "ve lezyon segmentasyonu yapan derin öğrenme tabanlı iki aşamalı bir "
    "sistem geliştirilmiştir. Birinci aşamada APTOS 2019 (3.662 görüntü) ve "
    "IDRiD (413 ek görüntü) veri setleri kullanılarak EfficientNet-B3 "
    "mimarisi ile beş sınıflı (No DR, Mild, Moderate, Severe, PDR) bir "
    "sınıflandırıcı eğitilmiştir. İkinci aşamada IDRiD veri seti üzerinde "
    "U-Net mimarisi ile dört farklı lezyon tipi (mikroanevrizma, kanama, "
    "sert eksuda ve yumuşak eksuda) piksel düzeyinde segmente edilmiştir. "
    "Ayrıca Grad-CAM++ yöntemiyle modelin karar verdiği bölgeler "
    "görselleştirilmiştir. Önerilen sınıflandırıcı APTOS validasyon setinde "
    "0.885 Cohen Quadratic Kappa skoru ve %80.76 doğruluk elde etmiştir. "
    "Segmentasyon modeli IDRiD test setinde ortalama 0.36 Dice skoru "
    "(mikroanevrizma hariç 0.49) başarı göstermiştir. Çalışmada ayrıca veri "
    "setleri arası domain shift sorunu ve donanım kısıtlamalarının "
    "performansa etkisi tartışılmıştır."
)
add_para(ozet, align='justify', space_after=18, line_spacing=1.5,
         indent_first=1.0)

add_para('Anahtar kelimeler:', bold=True, size=11, align='left',
         space_after=0, line_spacing=1.5)
add_para('Diyabetik Retinopati, Derin Öğrenme, Görüntü İşleme, '
         'Sınıflandırma, Segmentasyon, U-Net, EfficientNet, Grad-CAM, '
         'Fundus Görüntüsü, Açıklanabilir Yapay Zeka.',
         size=11, align='justify', space_after=12, line_spacing=1.5)

page_break()


# ====================== 1. GİRİŞ ======================
add_heading_main('1. GİRİŞ')

add_para(
    "Diyabetik retinopati (DR), diyabet hastalığında kan şekerinin uzun süre "
    "yüksek seyretmesi sonucunda retinal kan damarlarında meydana gelen "
    "hasarla karakterize edilen mikrovasküler bir komplikasyondur. "
    "Uluslararası Diyabet Federasyonu'nun 2021 yılı raporuna göre dünya "
    "genelinde yaklaşık 537 milyon yetişkin diyabet hastasıdır ve bu sayının "
    "2045 yılında 783 milyona ulaşması öngörülmektedir [1]. Diyabet "
    "hastalarının yaklaşık üçte birinde DR gelişmekte, bunların önemli bir "
    "kısmı görme tehdidi oluşturan ileri evreye ilerlemektedir [2]. Hastalık "
    "özellikle çalışma yaşındaki yetişkinler (20-65 yaş) arasında görme "
    "kaybının başlıca nedenlerinden biridir.",
    align='justify', indent_first=1.0)

add_para(
    "DR klinik olarak retinopati yok (No DR), hafif (Mild), orta "
    "(Moderate), şiddetli (Severe) ve proliferatif (PDR) olmak üzere beş "
    "kategoride sınıflandırılır [3]. Sınıflandırma; mikroanevrizma, kanama, "
    "sert eksuda ve yumuşak eksuda gibi lezyonların retinadaki yoğunluğu ve "
    "dağılımına dayanır. Klinik uygulamada tanı deneyimli oftalmologların "
    "fundus fotoğraflarını manuel olarak incelemesiyle konulmaktadır. Ancak "
    "diyabet prevalansının yüksekliği, uzman yetersizliği ve "
    "değerlendirici-arası tutarsızlık tarama programlarının etkinliğini "
    "sınırlandırmaktadır.",
    align='justify', indent_first=1.0)

add_para(
    "Son yıllarda derin öğrenme tabanlı bilgisayarlı tanı sistemleri DR "
    "tespitinde umut verici sonuçlar göstermektedir. Özellikle evrişimli "
    "sinir ağları (CNN) ile yapılan çalışmalarda oftalmolog seviyesine yakın "
    "doğruluk oranları rapor edilmektedir [4][5]. Ancak farklı kameralar ve "
    "görüntüleme protokolleri arasındaki domain shift problemi, modellerin "
    "gerçek klinik ortamlarda kullanılabilirliğini önemli ölçüde "
    "sınırlandırmaktadır.",
    align='justify', indent_first=1.0)

add_para(
    "Bu projede fundus görüntülerinden DR şiddet sınıflandırması ve lezyon "
    "bölge segmentasyonu yapan derin öğrenme tabanlı iki aşamalı bir sistem "
    "geliştirilmiştir. Önerilen sistem (i) EfficientNet-B3 mimarisi ile beş "
    "sınıflı DR şiddet sınıflandırması, (ii) U-Net mimarisi ile dört farklı "
    "lezyon tipinin (mikroanevrizma, kanama, sert/yumuşak eksuda) piksel "
    "düzeyinde segmentasyonu ve (iii) Grad-CAM++ yöntemiyle modelin karar "
    "mekanizmasının görselleştirilmesi olmak üzere üç temel bileşenden "
    "oluşmaktadır. APTOS 2019 ve IDRiD olmak üzere iki halka açık veri seti "
    "kullanılarak hem aynı-veri-seti hem de çapraz-veri-seti değerlendirmesi "
    "yapılmış, domain shift probleminin etkisi sayısal olarak ortaya "
    "konmuştur.",
    align='justify', indent_first=1.0)

add_para(
    "Bu raporun yapısı şöyledir: İkinci bölümde DR tespitine yönelik mevcut "
    "literatür özetlenmektedir. Üçüncü bölümde kullanılan veri setleri, ön "
    "işleme süreci, model mimarileri ve eğitim parametreleri "
    "detaylandırılmaktadır. Dördüncü bölümde sınıflandırma ve segmentasyon "
    "sonuçları sunulmakta; beşinci bölümde elde edilen bulgular tartışılmakta, "
    "sınırlılıklar ve gelecek çalışmalar değerlendirilmektedir. Altıncı bölümde "
    "ise çalışmanın genel sonuçları aktarılmaktadır.",
    align='justify', indent_first=1.0)

page_break()


# ====================== 2. LİTERATÜR TARAMASI ======================
add_heading_main('2. LİTERATÜR TARAMASI')

add_para(
    "Diyabetik retinopatinin otomatik tespitine yönelik çalışmalar 1990'lı "
    "yılların sonlarında başlamış, başlangıçta el yapımı öznitelikler ve "
    "klasik makine öğrenmesi yöntemleri (destek vektör makineleri, karar "
    "ağaçları, k-en yakın komşuluk) kullanılmıştır [6]. Ancak bu yöntemler "
    "lezyon özelliklerinin manuel olarak tanımlanmasını gerektirdiğinden "
    "esneklikleri sınırlı kalmıştır. 2015 yılında Google AI grubunun "
    "yayınladığı çalışmada Inception-v3 mimarisi 128.000 fundus görüntüsü "
    "üzerinde eğitilerek oftalmolog seviyesinde başarı gösterilmiş ve derin "
    "öğrenmenin DR tespitinde dönüm noktası olmuştur [4].",
    align='justify', indent_first=1.0)

add_para(
    "2019 yılında APTOS (Asia Pacific Tele-Ophthalmology Society) "
    "tarafından düzenlenen Kaggle yarışması, derin öğrenmeye dayalı DR "
    "sınıflandırma çalışmalarının yaygınlaşmasında önemli bir rol oynamıştır. "
    "Yarışmada ilk sıralara giren takımlar EfficientNet ve ResNet "
    "mimarilerini, ağırlıklı ortalama (weighted average) toplulaştırma "
    "yöntemleriyle birlikte kullanmış ve 0.93 üstü Cohen Quadratic Kappa "
    "skoru elde etmiştir [5]. Bu çalışmaların önemli bir bulgu olarak "
    "vurguladığı nokta ön işleme yöntemlerinin (özellikle Ben Graham'ın "
    "önerdiği fundus crop ve renk normalizasyonu) başarıya katkısının model "
    "mimarisi kadar önemli olduğudur.",
    align='justify', indent_first=1.0)

add_para(
    "Lezyon segmentasyonu konusunda en kapsamlı halka açık veri seti olan "
    "IDRiD (Indian Diabetic Retinopathy Image Dataset), 2018 yılında yayınlanmıştır "
    "ve dört farklı lezyon tipi (mikroanevrizma, kanama, sert ve yumuşak "
    "eksuda) ile optik disk için piksel düzeyinde maskeler içermektedir [7]. "
    "Bu veri seti üzerinde yapılan çalışmaların çoğu U-Net ve türevleri "
    "(U-Net++, Attention U-Net, U-Net+EfficientNet encoder) kullanmaktadır. "
    "Literatürde rapor edilen Dice skorları sert eksuda için 0.65-0.85, "
    "kanama için 0.45-0.70, mikroanevrizma için 0.35-0.55 aralıklarındadır. "
    "Mikroanevrizmalar (3-10 piksel boyutunda) tüm segmentasyon "
    "çalışmalarında en zorlu lezyon olarak rapor edilmektedir.",
    align='justify', indent_first=1.0)

add_para(
    "Açıklanabilirlik (explainability) konusunda Selvaraju ve arkadaşlarının "
    "2017 yılında yayınladığı Grad-CAM (Gradient-weighted Class Activation "
    "Mapping) yöntemi, CNN'lerin karar sürecini görselleştirmek için yaygın "
    "olarak kullanılmaktadır [8]. Bu yöntemin geliştirilmiş versiyonu olan "
    "Grad-CAM++ daha keskin ve doğruluk olarak daha güvenilir ısı haritaları "
    "üreterek medikal görüntü yorumlamada tercih edilmektedir.",
    align='justify', indent_first=1.0)

add_para(
    "Domain shift problemi son yıllarda yapay zeka destekli medikal "
    "uygulamaların gerçek klinik ortama aktarımı önünde en büyük engellerden "
    "biri olarak ortaya çıkmıştır. Farklı ülkelerden, farklı kameralarla "
    "(Canon CR-DGi, Topcon NW-400, Kowa VX-10) toplanmış veri setleri "
    "arasında performans düşüşleri %30-70 aralığında ölçümlenmiştir [9]. Bu "
    "problemin üzerine domain adaptation, domain generalization ve federated "
    "learning gibi yaklaşımlar önerilmektedir.",
    align='justify', indent_first=1.0)

page_break()


# ====================== 3. MATERYAL VE YÖNTEM ======================
add_heading_main('3. MATERYAL VE YÖNTEM')

add_heading_sub('3.1 Veri Setleri')
add_para(
    "Bu çalışmada iki halka açık fundus görüntü veri seti kullanılmıştır: "
    "APTOS 2019 Blindness Detection ve IDRiD (Indian Diabetic Retinopathy "
    "Image Dataset).",
    align='justify', indent_first=1.0)

add_para('APTOS 2019:', bold=True, size=12, align='left', space_after=4)
add_para(
    "APTOS 2019 Blindness Detection veri seti, Asia Pacific "
    "Tele-Ophthalmology Society tarafından düzenlenen Kaggle yarışması için "
    "yayınlanmıştır. Veri seti Hindistan'ın Aravind Göz Hastanesi'nde Topcon "
    "fundus kamerası ile çekilmiş 3.662 adet eğitim görüntüsünü içerir. Her "
    "görüntü deneyimli oftalmologlar tarafından 0-4 aralığında "
    "etiketlenmiştir (0: No DR, 1: Mild, 2: Moderate, 3: Severe, 4: PDR). "
    "Görüntü boyutları 640x480 ile 4288x2848 piksel arasında değişken olup, "
    "çalışmada kullanılan 100 örnek üzerinden hesaplanan ortalama boyut "
    "1986x1511 pikseldir.",
    align='justify', indent_first=1.0)

add_para('IDRiD:', bold=True, size=12, align='left', space_after=4)
add_para(
    "Indian Diabetic Retinopathy Image Dataset, Hindistan Teknoloji "
    "Enstitüsü tarafından 2018 yılında yayınlanmıştır ve iki ana alt veri "
    "setinden oluşmaktadır: B (Disease Grading) ve A (Segmentation). "
    "B alt veri seti 413 eğitim ve 103 test görüntüsü olmak üzere 516 adet "
    "sınıflandırma örneklerini içerir. A alt veri seti ise 54 eğitim ve "
    "27 test olmak üzere 81 adet piksel-seviyesinde maskelenmiş lezyon "
    "segmentasyonu görüntüsü içerir. Görüntüler 4288x2848 piksel sabit "
    "çözünürlüktedir.",
    align='justify', indent_first=1.0)

add_figure('fig01_class_distribution.png',
           'Şekil 1: Veri setlerinin sınıf bazlı dağılımı. Soldan sağa: '
           'APTOS, IDRiD-B (eğitim), IDRiD-B (test).')

add_heading_sub('3.2 Görüntü Ön İşleme')
add_para(
    "Fundus görüntüleri farklı boyutlarda ve aydınlatma koşullarında "
    "çekildiğinden doğrudan derin öğrenme modeline beslenmeye uygun "
    "değildir. Bu nedenle dört aşamalı bir ön işleme pipeline uygulanmıştır.",
    align='justify', indent_first=1.0)

add_para(
    "Birinci aşamada fundus crop işlemi uygulanmıştır. Bu yöntem Ben "
    "Graham'ın 2015 Kaggle DR yarışmasında önerdiği tekniktir. Yöntem, "
    "görüntüdeki siyah arka plan bölgelerini tespit eder ve yalnızca retinal "
    "daireyi koruyarak görüntüyü kırpar. Tespit, gri tonlamaya çevrilmiş "
    "görüntüdeki piksel yoğunluk değerinin belirli bir eşikten (tol=7) "
    "büyük olduğu pikselleri seçerek yapılır.",
    align='justify', indent_first=1.0)

add_para(
    "İkinci aşamada CLAHE (Contrast Limited Adaptive Histogram Equalization) "
    "yöntemi uygulanmıştır. CLAHE klasik histogram eşitlemeden farklı olarak "
    "görüntüyü küçük bölgelere (tile) ayırır ve her bölgede ayrı ayrı "
    "histogram eşitleme yapar; gürültüyü kontrol etmek için de bir kontrast "
    "sınırı (clipLimit) kullanır. Bu çalışmada LAB renk uzayında L (luminance) "
    "kanalına clipLimit=2.0 ve tileGridSize=(8,8) parametreleriyle CLAHE "
    "uygulanmıştır. Bu yöntem özellikle mikroanevrizma ve sert eksuda gibi "
    "lezyonların kontrastını arttırarak modelin lezyonları tespit etmesini "
    "kolaylaştırmaktadır.",
    align='justify', indent_first=1.0)

add_para(
    "Üçüncü aşamada görüntü, kare boyut korunarak siyah kenarlıkla "
    "doldurulmuş (zero-padded) ve son aşamada 384x384 piksele yeniden "
    "boyutlandırılmıştır. Bu boyut EfficientNet-B3 mimarisinin eğitim için "
    "standart girdi boyutudur ve ImageNet ortalama/standart sapma değerleri "
    "([0.485, 0.456, 0.406] / [0.229, 0.224, 0.225]) ile normalize "
    "edilmiştir.",
    align='justify', indent_first=1.0)

add_figure('fig02_preprocessing_pipeline.png',
           'Şekil 2: Fundus görüntü ön-işleme pipeline. Soldan sağa: '
           'orijinal görüntü, fundus crop, CLAHE uygulaması ve final 384x384 '
           'boyutlandırma aşamalarının üç farklı sınıf üzerindeki etkisi.')

add_heading_sub('3.3 Veri Bölümleme')
add_para(
    "APTOS veri seti stratified random sampling yöntemiyle %80 eğitim "
    "(2.929 görüntü) ve %20 validasyon (733 görüntü) olmak üzere ikiye "
    "ayrılmıştır. IDRiD B-Disease Grading eğitim seti (413 görüntü) "
    "eğitim setine eklenerek toplam 3.342 eğitim örneği elde edilmiştir. "
    "IDRiD-B test seti (103 görüntü) ise external test seti olarak ayrı "
    "tutulmuş ve modelin eğitim sırasında görmediği bağımsız bir veri "
    "üzerinde performansını ölçmek için kullanılmıştır. Bu yaklaşım "
    "sınıflandırma modelinin çapraz-veri-seti generalizasyon kapasitesini "
    "değerlendirmek için önemlidir.",
    align='justify', indent_first=1.0)

add_para(
    "Sınıf dağılımına bakıldığında veri setlerinde belirgin bir dengesizlik "
    "olduğu gözlemlenmektedir. APTOS train setinde No DR sınıfı %49.3 "
    "oranında iken Severe sınıfı yalnızca %5.3 oranında temsil edilmektedir. "
    "Bu dengesizliği gidermek için eğitimde class-weighted "
    "CrossEntropyLoss kullanılmıştır; her sınıf için ağırlık w_i = N / (K * "
    "n_i) formülü ile hesaplanmıştır (N: toplam örnek sayısı, K: sınıf "
    "sayısı, n_i: i. sınıftaki örnek sayısı). Bu yöntem azınlık sınıfların "
    "kayıp fonksiyonuna katkılarını arttırarak modelin tüm sınıfları öğrenmesini "
    "teşvik etmektedir.",
    align='justify', indent_first=1.0)

add_figure('fig03_augmentation_examples.png',
           'Şekil 3: Veri artırma (augmentation) örnekleri. Aynı görüntü '
           'farklı flip, döndürme, parlaklık ve bulanıklık kombinasyonlarıyla '
           'sekiz farklı varyasyonda gösterilmektedir.',
           width_cm=15)

add_heading_sub('3.4 Sınıflandırma Modeli')
add_para(
    "Sınıflandırma için EfficientNet-B3 mimarisi tercih edilmiştir. Tan ve "
    "Le tarafından 2019 yılında önerilen EfficientNet ailesi, klasik CNN "
    "mimarilerine kıyasla daha az parametre ile daha yüksek doğruluk "
    "sağlayan compound scaling yöntemiyle geliştirilmiştir [10]. B3 varyantı "
    "yaklaşık 10.7 milyon parametreye sahip olup, 384x384 girdi boyutu için "
    "optimize edilmiştir.",
    align='justify', indent_first=1.0)

add_para(
    "Model ImageNet veri seti üzerinde önceden eğitilmiş (pretrained) "
    "ağırlıklarla başlatılmış ve transfer learning yöntemi ile DR tanı "
    "problemine adapte edilmiştir. Son katmandaki fully-connected katman "
    "5 sınıf çıkışına göre yeniden tanımlanmış ve dropout rate 0.3 olarak "
    "ayarlanarak overfitting önlenmeye çalışılmıştır.",
    align='justify', indent_first=1.0)

add_heading_sub('3.5 Segmentasyon Modeli')
add_para(
    "Lezyon segmentasyonu için U-Net mimarisi tercih edilmiştir. "
    "Ronneberger ve arkadaşlarının 2015 yılında biyomedikal görüntü "
    "segmentasyonu için önerdiği bu mimari, encoder-decoder yapısında "
    "skip-connection mekanizması kullanarak hem yüksek seviyeli semantik "
    "özellikleri hem de düşük seviyeli mekansal bilgileri korumaktadır [11]. "
    "Bu çalışmada encoder olarak EfficientNet-B0 (yaklaşık 4 milyon "
    "parametre, ImageNet pretrained) kullanılmış ve toplam model "
    "parametre sayısı 6.25 milyon olmuştur.",
    align='justify', indent_first=1.0)

add_para(
    "Çıkış katmanı 4 kanal olup her kanal bir lezyon tipini (MA, HE, EX, "
    "SE) temsil etmektedir. Bu sayede tek bir model ile multi-label "
    "segmentasyon yapılabilmektedir. Çıkış aktivasyon fonksiyonu sigmoid "
    "olup, her piksel için lezyon olma olasılığı 0-1 aralığında "
    "üretilmektedir; 0.5 eşiği ile binary maskelere dönüştürülmüştür.",
    align='justify', indent_first=1.0)

add_para(
    "Kayıp fonksiyonu olarak BCE (Binary Cross-Entropy) ve Dice loss "
    "kombinasyonu (eşit ağırlıklı, w_BCE = w_Dice = 0.5) kullanılmıştır. "
    "Bu kombinasyon küçük ve seyrek lezyonların (özellikle "
    "mikroanevrizmaların) başarılı segmentasyonu için literatürde yaygın "
    "kullanılmaktadır.",
    align='justify', indent_first=1.0)

add_heading_sub('3.6 Açıklanabilirlik (Grad-CAM)')
add_para(
    "Modelin karar mekanizmasını görselleştirmek için Grad-CAM++ yöntemi "
    "kullanılmıştır. Bu yöntem CNN'nin son evrişim katmanında oluşan "
    "özellik haritalarının tahmin edilen sınıftan gelen gradyanlarla "
    "ağırlıklandırılması suretiyle bir ısı haritası (heatmap) üretmektedir. "
    "Isı haritası yüksek değer alan bölgeler modelin tahminine en çok "
    "katkı sağlayan görüntü bölgelerini göstermektedir. Klinik bağlamda "
    "bu yöntem modelin lezyona mı yoksa başka bir yapıya mı (optik disk, "
    "damar) baktığını doğrulamada kullanılmaktadır.",
    align='justify', indent_first=1.0)

add_heading_sub('3.7 Eğitim Detayları')
add_para(
    "Tüm eğitimler Google Colab platformunda NVIDIA Tesla T4 GPU "
    "(16 GB VRAM) üstünde PyTorch 2.5.1 kütüphanesi kullanılarak "
    "gerçekleştirilmiştir. Mixed precision (AMP) eğitim kullanılarak "
    "GPU bellek tüketimi yarıya indirilmiş ve eğitim süresi yaklaşık "
    "iki kat hızlanmıştır.",
    align='justify', indent_first=1.0)

add_table(
    headers=['Hiperparametre', 'Sınıflandırma', 'Segmentasyon'],
    rows=[
        ('Model', 'EfficientNet-B3', 'U-Net + EfficientNet-B0'),
        ('Parametre Sayısı', '10.7M', '6.25M'),
        ('Girdi Boyutu', '384x384', '512x512'),
        ('Batch Size', '32', '4'),
        ('Epoch Sayısı', '20', '60'),
        ('Optimizer', 'AdamW', 'AdamW'),
        ('Başlangıç LR', '3e-4', '1e-4'),
        ('Weight Decay', '1e-4', '1e-5'),
        ('Scheduler', 'Cosine + Warmup (2 ep)', 'Cosine + Warmup (3 ep)'),
        ('Kayıp Fonksiyonu', 'CrossEntropy (weighted)', 'BCE + Dice (0.5 / 0.5)'),
        ('Label Smoothing', '0.05', '-'),
        ('Eğitim Süresi', '~21 dakika', '~63 dakika'),
    ],
    col_widths_cm=[6, 5, 5]
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Tablo 1: Eğitim hiperparametreleri.')
run.bold = True
run.font.size = Pt(11)

page_break()


# ====================== 4. BULGULAR ======================
add_heading_main('4. BULGULAR')

add_heading_sub('4.1 Sınıflandırma Sonuçları')
add_para(
    "EfficientNet-B3 modeli APTOS validasyon seti üzerinde 20 epoch boyunca "
    "eğitilmiştir. Şekil 4'te gösterilen eğitim eğrilerinde train ve "
    "validation kayıpları yakın seyretmekte ve overfitting belirtisi "
    "gözlemlenmemektedir. Cohen Quadratic Kappa skoru ilk 5 epoch'ta "
    "0.85 seviyesine yükselmiş, sonraki epoch'larda yavaşça artarak en iyi "
    "değer olan 0.885'i epoch 19'da elde etmiştir.",
    align='justify', indent_first=1.0)

add_figure('fig04_training_curves.png',
           'Şekil 4: Sınıflandırma modeli eğitim eğrileri. Sol üst: train/val '
           'kayıp; sağ üst: train/val accuracy; sol alt: Quadratic Kappa ve '
           'F1-macro; sağ alt: warmup+cosine learning rate programı.',
           width_cm=15)

add_para(
    "Tablo 2'de APTOS validasyon ve IDRiD external test setleri üzerindeki "
    "ana metrikler özetlenmiştir. Validasyon setinde model 0.885 Kappa, "
    "%80.76 doğruluk ve 0.656 F1-macro skoru elde etmiştir. Bu skorlar "
    "APTOS 2019 Kaggle yarışmasının üst %10'una denk gelen seviyededir.",
    align='justify', indent_first=1.0)

add_table(
    headers=['Veri Seti', 'Accuracy', 'F1-Macro', 'Kappa (quad.)'],
    rows=[
        ('APTOS Validasyon', '0.808', '0.656', '0.885'),
        ('IDRiD External Test', '0.214', '0.156', '0.146'),
    ],
    col_widths_cm=[5, 3.5, 3.5, 3.5]
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Tablo 2: Sınıflandırma performansı (Internal Validation vs External Test).')
run.bold = True
run.font.size = Pt(11)

add_para(
    "Karışıklık matrisi (Şekil 5) detaylı incelendiğinde, APTOS "
    "validasyon setinde No DR sınıfı %97 doğruluk ile en iyi performanslı "
    "sınıf olmuştur. Mild ve Severe sınıflarda hataların çoğu komşu sınıflar "
    "(Moderate, PDR) ile karıştırma yönündedir; bu durum klinik olarak "
    "kabul edilebilir, çünkü komşu sınıflarda lezyon dağılımları benzer "
    "özellikler taşımaktadır.",
    align='justify', indent_first=1.0)

add_figure('fig05_confusion_matrices.png',
           'Şekil 5: APTOS validasyon ve IDRiD external test setleri için '
           'normalize edilmiş karışıklık matrisleri. Renk yoğunluğu artan '
           'normalize oranında koyulaşmaktadır.',
           width_cm=15)

add_para(
    "Ancak IDRiD external test setinde belirgin bir performans düşüşü "
    "gözlemlenmiştir (Kappa 0.146, accuracy %21.4). Karışıklık matrisinde "
    "modelin No DR sınıfındaki 34 örnekten hiçbirini doğru tahmin "
    "edemediği görülmektedir; modelin tahminleri Moderate, Severe ve PDR "
    "sınıflarına dağılmaktadır. Bu sonuç, modelin APTOS görüntülerinin "
    "özelliklerine aşırı uyduğunu (over-specialization) ve IDRiD'in farklı "
    "kamera/aydınlatma koşulları altında çekilmiş görüntülerine adapte "
    "olamadığını göstermektedir.",
    align='justify', indent_first=1.0)

add_heading_sub('4.2 Domain Adaptation Sonuçları')
add_para(
    "External test setindeki düşük performansı iyileştirmek amacıyla "
    "model IDRiD-B eğitim seti (413 görüntü) üzerinde 5 epoch boyunca "
    "düşük öğrenme oranıyla (1e-5) fine-tune edilmiştir. Bu yaklaşım "
    "modelin önceden öğrendiği genel özellikleri korurken IDRiD görüntü "
    "özelliklerine adapte olmasını hedeflemektedir.",
    align='justify', indent_first=1.0)

add_figure('fig06_confusion_finetuned.png',
           'Şekil 6: IDRiD-B üzerinde fine-tune sonrası karışıklık matrisleri. '
           'Solda APTOS validasyon, sağda IDRiD external test sonuçları.',
           width_cm=15)

add_para(
    "Şekil 6 ve Şekil 7'de görülebileceği üzere fine-tune sonrasında "
    "IDRiD test Kappa skoru 0.146'dan 0.153'e çıkmış, ancak APTOS "
    "validasyon Kappa skoru 0.885'ten 0.867'ye düşmüştür. Bu marjinal "
    "iyileşme, mevcut IDRiD-B eğitim setinin (413 görüntü) çapraz-veri-seti "
    "adaptasyonu için yetersiz olduğunu göstermektedir. Daha büyük hedef "
    "veri seti veya domain adaptation tekniklerinin (adversarial training, "
    "MMD loss) kullanılması gerekmektedir.",
    align='justify', indent_first=1.0)

add_figure('fig07_finetune_effect.png',
           'Şekil 7: Fine-tune süreci boyunca APTOS Val ve IDRiD Test '
           'Quadratic Kappa skorlarının gelişimi. Epoch 0 fine-tune öncesi, '
           '1-5 fine-tune sırasındaki epochlardır.',
           width_cm=14)

add_heading_sub('4.3 Segmentasyon Sonuçları')
add_para(
    "U-Net modeli IDRiD segmentasyon eğitim seti (54 görüntü) üzerinde "
    "60 epoch boyunca eğitilmiş ve IDRiD test seti (27 görüntü) üzerinde "
    "değerlendirilmiştir. Eğitim ve test kayıpları yakın seyretmekte, "
    "overfitting belirtisi gözlemlenmemektedir (Şekil 8). Bu durum modelin "
    "kapasitesinin yeterli olduğunu ve eğitim verisinin sınırlı olmasının "
    "asıl kısıtlama olduğunu göstermektedir.",
    align='justify', indent_first=1.0)

add_figure('fig_seg_training_curves.png',
           'Şekil 8: Segmentasyon modeli eğitim eğrileri. Solda train/test '
           'kaybı (BCE + Dice), sağda lezyon bazında test Dice skorları.',
           width_cm=15)

add_para(
    "Tablo 3 ve Şekil 9'da lezyon bazında final test metrikleri "
    "sunulmuştur. Sert eksuda (EX) ve kanama (HE) için sırasıyla 0.593 "
    "ve 0.463 Dice skorları elde edilmiştir; bu değerler literatürde "
    "rapor edilen ortalamalara yakındır. Yumuşak eksuda (SE) için 0.402 "
    "Dice skoru, sınırlı eğitim verisi (yalnızca 26 maskelenmiş görüntü) "
    "dikkate alındığında kabul edilebilir bir başarıdır. Mikroanevrizma "
    "(MA) için Dice skoru 0.000 olarak rapor edilmiştir; bu durum "
    "donanım/çözünürlük kısıtlaması ile açıklanmaktadır (5.3. bölüme bakınız).",
    align='justify', indent_first=1.0)

add_table(
    headers=['Lezyon', 'Dice', 'IoU', 'Maskelenmiş Görüntü (eğitim)'],
    rows=[
        ('Mikroanevrizma (MA)', '0.000', '0.000', '54'),
        ('Kanama (HE)', '0.463', '0.301', '53'),
        ('Sert Eksuda (EX)', '0.593', '0.422', '54'),
        ('Yumuşak Eksuda (SE)', '0.402', '0.252', '26'),
        ('Ortalama (4 sınıf)', '0.365', '0.244', '-'),
        ('Ortalama (MA hariç)', '0.486', '0.325', '-'),
    ],
    col_widths_cm=[5, 2.5, 2.5, 4.5]
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Tablo 3: Segmentasyon performansı (IDRiD Test seti, 27 görüntü).')
run.bold = True
run.font.size = Pt(11)

add_figure('fig_seg_per_class_metrics.png',
           'Şekil 9: Lezyon bazında final Dice ve IoU skorları.',
           width_cm=13)

add_para(
    "Şekil 10'da test seti üzerinden seçilen üç örneklem için ground "
    "truth ve model tahminleri yan yana sunulmuştur. Sert eksudaların "
    "(sarı) modelle doğruya yakın şekilde tespit edildiği, kanama "
    "(turuncu) tahminlerinin ise GT'ye göre biraz daha agresif (geniş "
    "alanlı) yapıldığı gözlemlenmektedir.",
    align='justify', indent_first=1.0)

add_figure('fig_seg_predictions.png',
           'Şekil 10: Segmentasyon sonuçları. Soldan sağa: orijinal görüntü, '
           'ground truth lezyon maskeleri ve model tahminleri. Renkler: '
           'kırmızı=MA, turuncu=HE, sarı=EX, mavi=SE.',
           width_cm=15)

add_heading_sub('4.4 Grad-CAM Analizi')
add_para(
    "Şekil 11'de beş farklı DR sınıfından birer örnek için Grad-CAM++ "
    "ısı haritaları sunulmuştur. Modelin No DR sınıfında dağıt bir "
    "aktivasyon paterni gösterdiği, lezyon içeren sınıflarda ise (özellikle "
    "Moderate, Severe ve PDR) ısı haritasının gerçek lezyon bölgelerinde "
    "(sert eksudalar, mikroanevrizma kümeleri, hemorajik alanlar) "
    "yoğunlaştığı gözlemlenmiştir. Bu sonuç, modelin karar verme sürecinde "
    "klinik olarak ilgili lezyon bölgelerine baktığını ve karar "
    "güvenirliğini doğrulamaktadır.",
    align='justify', indent_first=1.0)

add_figure('fig_gradcam.png',
           'Şekil 11: Beş farklı DR sınıfından seçilen örnek görüntüler '
           'için Grad-CAM++ ısı haritaları. Solda orijinal görüntü, sağda '
           'ısı haritası overlay.',
           width_cm=14)

page_break()


# ====================== 5. TARTIŞMA VE SINIRLILIKLAR ======================
add_heading_main('5. TARTIŞMA VE SINIRLILIKLAR')

add_heading_sub('5.1 Sınıf Dengesizliği ve Eğitim Stratejisi')
add_para(
    "Veri setlerindeki belirgin sınıf dengesizliği (No DR %49.3, Severe "
    "%5.3) modeli dengeli öğrenmeye zorlamak için ağırlıklı kayıp "
    "fonksiyonu (class-weighted CrossEntropy) ve label smoothing (0.05) "
    "kombinasyonu uygulanmıştır. Sonuçlarda azınlık sınıflarında da makul "
    "performans elde edilmesi bu yaklaşımın etkili olduğunu göstermektedir. "
    "Yine de Severe sınıfında F1 skoru 0.41 olup, sınıfın az sayıda örnek "
    "içermesi nedeniyle bu sınıfta daha fazla iyileştirme alanı bulunmaktadır. "
    "Gelecekteki çalışmalarda focal loss veya MixUp/CutMix gibi gelişmiş "
    "augmentation yöntemleri denenebilir.",
    align='justify', indent_first=1.0)

add_heading_sub('5.2 Domain Shift Problemi')
add_para(
    "Bu çalışmanın en önemli bulgusu modelin APTOS validasyon setindeki "
    "yüksek başarısının (Kappa 0.885) IDRiD external test setinde "
    "dramatik bir şekilde düştüğüdür (Kappa 0.146). Bu fark, derin "
    "öğrenme modellerinin klinik kullanıma sokulmadan önce çoklu-merkez "
    "veri setleri üzerinde valide edilmesi gerektiğini açıkça "
    "göstermektedir. APTOS görüntüleri Hindistan'ın Aravind Göz "
    "Hastanesi'nde Topcon NW-400 kamerası ile, IDRiD görüntüleri ise "
    "aynı ülkenin farklı bir kliniğinde Kowa VX-10α kamerası (50 "
    "derece görüş alanlı) ile çekilmiştir. Bu farklılıklar; aydınlatma "
    "düzeyi, renk dengesi, optik distorsiyon ve görüntü kalitesi gibi "
    "faktörler üzerinde belirgin etkiler yaratmaktadır.",
    align='justify', indent_first=1.0)

add_para(
    "Fine-tune denemesi (IDRiD-B üzerinde 5 epoch) IDRiD test "
    "performansını marjinal olarak iyileştirmiş (Kappa 0.146 → 0.153) "
    "ancak APTOS validasyon performansını biraz düşürmüştür. Bu sonuç, "
    "domain adaptation için daha gelişmiş tekniklerin (adversarial "
    "domain adaptation, MMD loss, batch normalization istatistik adaptasyonu, "
    "self-supervised pretraining) ve daha çeşitli hedef veri setlerinin "
    "gerekliliğini göstermektedir.",
    align='justify', indent_first=1.0)

add_heading_sub('5.3 Mikroanevrizma Segmentasyonunda Çözünürlük Sınırlılığı')
add_para(
    "Mikroanevrizmalar (MA) gerçek fundus görüntülerinde 3-10 piksel "
    "çapında küçük noktalardır. Bu çalışmada kullanılan 4288x2848 piksel "
    "boyutundaki IDRiD görüntüleri segmentasyon için 512x512 boyutuna "
    "küçültülmüştür. Bu küçültme işlemi sırasında mikroanevrizmalar "
    "0.5-1.2 piksel boyutuna inerek pratikte arka planla birleştirilmiştir. "
    "Sonuç olarak U-Net modeli MA için güvenli tahmin olarak sürekli arka "
    "plan tahmini üreterek Dice skorunu 0.000 dolayında tutmuştur.",
    align='justify', indent_first=1.0)

add_para(
    "Bu probleme literatürde önerilen çözüm patch-tabanlı eğitim "
    "yaklaşımıdır. Bu yaklaşımda yüksek çözünürlüklü (1024x1024 veya "
    "üstü) görüntülerden 256x256 veya 512x512 boyutunda örtüşen yamalar "
    "(patches) çıkarılır ve eğitim/çıkaresim bu yamalar üzerinde "
    "yapılır. Böylece mikroanevrizmaların piksel boyutu korunmuş olur. "
    "Ancak bu yaklaşım eğitim veri sayısını yaklaşık 50 kat arttırır, "
    "eğitim süresini ve donanım ihtiyacını büyük ölçüde büyütür. "
    "Mevcut donanım koşulları altında bu yöntem uygulanamamıştır.",
    align='justify', indent_first=1.0)

add_heading_sub('5.4 Donanımsal Kısıtlamalar ve Maliyet Analizi')
add_para(
    "Bu çalışmada Google Colab tarafından ücretsiz sağlanan NVIDIA "
    "Tesla T4 GPU (16 GB VRAM) kullanılmıştır. Colab'ın ücretsiz "
    "sürümü; 12 saatlik oturum süresi sınırlaması, kullanılmayan "
    "oturumlarda otomatik bağlantı kesilmesi ve sınırlı disk I/O hızı "
    "gibi kısıtlamalara sahiptir. Eğitim sırasında tekrarlayan oturum "
    "kayıpları yaşanmış, model ağırlıkları kalıcı depolamaya (Google "
    "Drive) periyodik olarak yedeklenmek zorunda kalınmıştır. Drive'dan "
    "her epoch'ta görüntü okumanın yavaşlığı nedeniyle ön-işlenmiş "
    "görüntüler yerel diske önbelleklenerek eğitim süresi yaklaşık 24 "
    "kat hızlandırılabilmiştir.",
    align='justify', indent_first=1.0)

add_para(
    "Mikroanevrizma segmentasyonu için gerekli olan yüksek çözünürlüklü "
    "(≥1024x1024) patch-tabanlı eğitim, daha yüksek VRAM'li GPU'lar "
    "(NVIDIA A100 40 GB veya RTX A6000 48 GB) gerektirmektedir. Bu tip "
    "GPU'lar bulut hizmetlerinde saat başına yaklaşık 2-5 USD "
    "ücretlendirilmekte; tam bir eğitim çevrimi (20-30 saat) için "
    "60-150 USD maliyet ortaya çıkmaktadır. Lokal donanım yatırımı "
    "tercih edildiğinde RTX A6000 GPU'nun fiyatı 4.500-5.500 USD, "
    "A100 GPU'nun ise 10.000-15.000 USD bandındadır. Bu durum, ücretsiz "
    "bulut altyapıları ile akademik bir lisans projesi kapsamında elde "
    "edilebilecek sonuçların doğal sınırını ortaya koymaktadır.",
    align='justify', indent_first=1.0)

add_para(
    "Özet olarak bu çalışmanın sonuçları, ücretsiz bulut kaynakları "
    "ile elde edilebilecek tipik bir performans seviyesini temsil "
    "etmekte; klinik düzeyde lezyon segmentasyonu için yeterli donanım "
    "yatırımının zorunlu olduğunu göstermektedir. Akademik çalışmalarda "
    "bu tür kısıtlamaların açıkça rapor edilmesi sonuçların gerçekçi "
    "olarak yorumlanmasını sağlamaktadır.",
    align='justify', indent_first=1.0)

page_break()


# ====================== 6. SONUÇ ======================
add_heading_main('6. SONUÇ')

add_para(
    "Bu çalışmada fundus görüntülerinden diyabetik retinopati şiddet "
    "sınıflandırması ve lezyon segmentasyonu yapan derin öğrenme tabanlı "
    "iki aşamalı bir sistem geliştirilmiştir. EfficientNet-B3 mimarisi "
    "ile eğitilen sınıflandırıcı, APTOS validasyon setinde 0.885 Cohen "
    "Quadratic Kappa skoru elde ederek literatürdeki üst seviye "
    "performanslara yakın bir başarı göstermiştir. U-Net mimarisi ile "
    "eğitilen segmentasyon modeli sert eksuda ve kanama için literatür "
    "ortalamalarına yakın (sırasıyla 0.593 ve 0.463 Dice) skorlar elde "
    "etmiştir.",
    align='justify', indent_first=1.0)

add_para(
    "Çalışmanın bir diğer önemli katkısı modelin karar sürecinin "
    "Grad-CAM++ yöntemiyle görselleştirilmesi ve modelin klinik olarak "
    "anlamlı lezyon bölgelerine baktığının doğrulanmasıdır. Bu, gerek "
    "akademik gerek klinik bağlamda yapay zeka destekli tanı sistemlerinin "
    "güvenilirliği açısından önemlidir.",
    align='justify', indent_first=1.0)

add_para(
    "Öte yandan çalışma; (i) APTOS-IDRiD arasında gözlemlenen domain shift "
    "problemi, (ii) mikroanevrizmaların düşük çözünürlükte segmente "
    "edilememesi ve (iii) ücretsiz bulut altyapılarının donanım "
    "kısıtlamaları gibi önemli sınırlılıklara sahiptir. Bu sınırlılıklar "
    "akademik dürüstlük ilkesi gereği açıkça rapor edilmiştir.",
    align='justify', indent_first=1.0)

add_para(
    "Gelecek çalışmalarda domain adaptation tekniklerinin uygulanması, "
    "patch-tabanlı yüksek çözünürlüklü segmentasyon yaklaşımlarının "
    "denenmesi ve çoklu-merkez veri setleriyle validasyon yapılması "
    "planlanmaktadır. Ek olarak modelin OCT (Optik Koherens Tomografi) "
    "görüntüleriyle birlikte kullanıldığı multi-modal yaklaşımlar; "
    "tanı başarısını önemli ölçüde arttırma potansiyeline sahiptir.",
    align='justify', indent_first=1.0)

add_para(
    "Bu projenin kaynak kodu (Python notebooks ve model dosyaları) "
    "GitHub üzerinde açık erişimle paylaşılmıştır: "
    "https://github.com/cantekinn/DR-Analysis-System",
    align='justify', indent_first=1.0)

page_break()


# ====================== KAYNAKLAR ======================
add_heading_main('KAYNAKLAR')

references = [
    ('[1]', 'International Diabetes Federation (2021). IDF Diabetes Atlas, '
     '10th edition. Brussels, Belgium.'),
    ('[2]', 'Wong, T. Y., Cheung, C. M. G., Larsen, M., Sharma, S., '
     '& Simo, R. (2016). Diabetic retinopathy. Nature Reviews Disease '
     'Primers, 2(1), 16012.'),
    ('[3]', 'Wilkinson, C. P., Ferris, F. L., Klein, R. E., Lee, P. P., '
     'Agardh, C. D., Davis, M., ... & Verdaguer, J. T. (2003). Proposed '
     'international clinical diabetic retinopathy and diabetic macular '
     'edema disease severity scales. Ophthalmology, 110(9), 1677-1682.'),
    ('[4]', 'Gulshan, V., Peng, L., Coram, M., Stumpe, M. C., Wu, D., '
     'Narayanaswamy, A., ... & Webster, D. R. (2016). Development and '
     'validation of a deep learning algorithm for detection of diabetic '
     'retinopathy in retinal fundus photographs. JAMA, 316(22), 2402-2410.'),
    ('[5]', 'Asia Pacific Tele-Ophthalmology Society (2019). APTOS 2019 '
     'Blindness Detection. Kaggle. '
     'https://www.kaggle.com/competitions/aptos2019-blindness-detection'),
    ('[6]', 'Acharya, U. R., Lim, C. M., Ng, E. Y. K., Chee, C., '
     '& Tamura, T. (2009). Computer-based detection of diabetes '
     'retinopathy stages using digital fundus images. Proceedings of the '
     'Institution of Mechanical Engineers, Part H: Journal of Engineering '
     'in Medicine, 223(5), 545-553.'),
    ('[7]', 'Porwal, P., Pachade, S., Kamble, R., Kokare, M., '
     'Deshmukh, G., Sahasrabuddhe, V., & Meriaudeau, F. (2018). Indian '
     'diabetic retinopathy image dataset (IDRiD): A database for diabetic '
     'retinopathy screening research. Data, 3(3), 25.'),
    ('[8]', 'Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., '
     'Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from '
     'deep networks via gradient-based localization. Proceedings of the '
     'IEEE International Conference on Computer Vision (ICCV), 618-626.'),
    ('[9]', 'Yang, Y., Shang, F., Wu, B., Yang, D., Wang, L., Xu, Y., '
     '... & Hu, B. (2021). Robust collaborative learning of patch-level and '
     'image-level annotations for diabetic retinopathy grading from fundus '
     'image. IEEE Transactions on Cybernetics, 52(11), 11407-11417.'),
    ('[10]', 'Tan, M., & Le, Q. V. (2019). EfficientNet: Rethinking model '
     'scaling for convolutional neural networks. International Conference '
     'on Machine Learning (ICML), 6105-6114.'),
    ('[11]', 'Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: '
     'Convolutional networks for biomedical image segmentation. Medical '
     'Image Computing and Computer-Assisted Intervention (MICCAI), '
     '234-241.'),
    ('[12]', 'Iakubovskii, P. (2019). Segmentation Models PyTorch. '
     'GitHub repository. '
     'https://github.com/qubvel/segmentation_models.pytorch'),
    ('[13]', 'Buslaev, A., Iglovikov, V. I., Khvedchenya, E., Parinov, A., '
     'Druzhinin, M., & Kalinin, A. A. (2020). Albumentations: Fast and '
     'flexible image augmentations. Information, 11(2), 125.'),
]

for kod, ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r1 = p.add_run(f'{kod} ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(ref)
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'


# ====================== KAYDET ======================
doc.save(str(OUTPUT))
print(f"Rapor olusturuldu: {OUTPUT}")
print(f"Toplam paragraf: {len(doc.paragraphs)}")
print()
print("Figur klasoru:", FIG_DIR)
if not FIG_DIR.exists():
    print("UYARI: Figur klasoru bulunamadi!")
else:
    pngs = list(FIG_DIR.glob('*.png'))
    print(f"Bulunan PNG sayisi: {len(pngs)}")
